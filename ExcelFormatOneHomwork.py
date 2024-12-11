# coding=utf-8
import xlrd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import RGBColor

# 声明一个word对象

doc = Document()

# 设置字体样式

doc.styles['Normal'].font.name = u'宋体'

doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

workbook = xlrd.open_workbook(r'C:\Users\Danny\Desktop\333.xls')  # 打开Excel
sheet = workbook.sheets()[0]
rows = sheet.nrows

print(type(sheet))

temp = '题型'
eum = ['A', 'B', 'C', 'D', 'E', 'F']

for i in range(rows):
    if i == 0:
        continue
    row = sheet.row(i)
    if temp != row[0].value:
        temp = row[0].value
        print('\t' + temp)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(temp)
        font = run.font
        font.size = Pt(24)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if row[0].value == '判断':
        # 题目
        print(str(i) + '：' + row[1].value)
        timuParagraph = doc.add_paragraph()
        timurun = timuParagraph.add_run(str(i) + '：' + row[1].value)
        timufont = timurun.font
        timufont.color.rgb = RGBColor(5, 145, 206)

        # 答案
        print('\t标准答案：' + row[2].value + '\t难易程度：' + row[10].value)
        daAnParagraph = doc.add_paragraph()
        daAnrun = daAnParagraph.add_run('\t标准答案：' + row[2].value + '\t难易程度：' + row[10].value)
        daAnrun = daAnrun.font
        daAnrun.color.rgb = RGBColor(221, 0, 27)

    else:
        # 题目
        print(str(i) + '：' + row[1].value)
        timuParagraph = doc.add_paragraph()
        timurun = timuParagraph.add_run(str(i) + '：' + row[1].value)
        timufont = timurun.font
        timufont.color.rgb = RGBColor(5, 145, 206)

        # 选项
        daAn = ''
        for n in range(6):
            if row[n + 2].value == '':
                continue
            daAn = daAn + '\t\t' + eum[n] + ':' + str(row[n + 2].value)
            if n % 2 == 1:
                daAn = daAn + '\n'
        if daAn[-1:] == '\n':
            daAn = daAn[:-1]
        print(daAn)
        xuanxiangParagraph = doc.add_paragraph(daAn)

        # 答案
        print('\t标准答案：' + row[8].value + '\t难易程度：' + row[10].value)
        daAnParagraph = doc.add_paragraph()
        daAnrun = daAnParagraph.add_run('\t标准答案：' + row[8].value + '\t难易程度：' + row[10].value)
        daAnrun = daAnrun.font
        daAnrun.color.rgb = RGBColor(221, 0, 27)
        # print(sheet.row(i))
        # sheet.cell[row = i,colum=1]

doc.save('111.docx')
